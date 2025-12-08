"""
Report Exporter Service
=======================
Exports reports to Word, Excel, PDF, and CSV formats.
"""

import io
import uuid
from decimal import Decimal
from datetime import datetime, timedelta
from typing import Dict, Optional, Any, Tuple
from django.core.files.base import ContentFile
from django.utils import timezone

from ..models import Report, ReportExport, ReportStatus, ExportFormat, ReportType


class ReportExporterService:
    """
    Service for exporting reports to various formats.
    Supports Excel, Word, PDF, and CSV.
    """
    
    # MIME types for export formats
    MIME_TYPES = {
        ExportFormat.EXCEL.value: 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        ExportFormat.WORD.value: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        ExportFormat.PDF.value: 'application/pdf',
        ExportFormat.CSV.value: 'text/csv',
        ExportFormat.JSON.value: 'application/json',
    }
    
    FILE_EXTENSIONS = {
        ExportFormat.EXCEL.value: '.xlsx',
        ExportFormat.WORD.value: '.docx',
        ExportFormat.PDF.value: '.pdf',
        ExportFormat.CSV.value: '.csv',
        ExportFormat.JSON.value: '.json',
    }
    
    def __init__(self):
        pass
    
    def export_report(
        self,
        report: Report,
        export_format: str,
        user,
        export_config: Optional[Dict] = None
    ) -> ReportExport:
        """
        Export a report to the specified format.
        Returns a ReportExport record with the generated file.
        """
        config = export_config or {}
        
        # Create export record
        export = ReportExport.objects.create(
            tenant_id=report.tenant_id,
            report=report,
            export_format=export_format,
            file_name=self._generate_filename(report, export_format),
            mime_type=self.MIME_TYPES.get(export_format, 'application/octet-stream'),
            status=ReportStatus.GENERATING.value,
            export_config=config,
            expires_at=timezone.now() + timedelta(days=7),  # Files expire after 7 days
            exported_by=user
        )
        
        try:
            # Generate the file based on format
            file_content, file_size = self._generate_file(
                report=report,
                export_format=export_format,
                config=config
            )
            
            # Save file
            export.file.save(
                export.file_name,
                ContentFile(file_content),
                save=False
            )
            export.file_size = file_size
            export.status = ReportStatus.COMPLETED.value
            export.save()
            
        except Exception as e:
            export.status = ReportStatus.FAILED.value
            export.error_message = str(e)
            export.save()
            raise
        
        return export
    
    def _generate_filename(self, report: Report, export_format: str) -> str:
        """Generate a filename for the export"""
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        extension = self.FILE_EXTENSIONS.get(export_format, '')
        safe_name = report.name.replace(' ', '_').replace('/', '-')[:50]
        return f"{safe_name}_{report.report_number}_{timestamp}{extension}"
    
    def _generate_file(
        self,
        report: Report,
        export_format: str,
        config: Dict
    ) -> Tuple[bytes, int]:
        """Generate file content based on format"""
        generators = {
            ExportFormat.EXCEL.value: self._generate_excel,
            ExportFormat.WORD.value: self._generate_word,
            ExportFormat.CSV.value: self._generate_csv,
            ExportFormat.JSON.value: self._generate_json,
            ExportFormat.PDF.value: self._generate_pdf,
        }
        
        generator = generators.get(export_format)
        if not generator:
            raise ValueError(f"Unsupported export format: {export_format}")
        
        return generator(report, config)
    
    # =================================================================
    # Excel Export
    # =================================================================
    
    def _generate_excel(self, report: Report, config: Dict) -> Tuple[bytes, int]:
        """Generate Excel file from report data"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl is required for Excel export. Install with: pip install openpyxl")
        
        wb = Workbook()
        ws = wb.active
        ws.title = report.report_type.replace('_', ' ').title()[:31]
        
        # Styles
        header_font = Font(bold=True, size=14)
        subheader_font = Font(bold=True, size=11)
        money_format = '#,##0.00'
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        header_fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        
        row = 1
        
        # Title
        ws.cell(row=row, column=1, value=report.name).font = header_font
        row += 1
        ws.cell(row=row, column=1, value=f"Period: {report.period_display}")
        row += 1
        ws.cell(row=row, column=1, value=f"Generated: {report.generation_completed_at.strftime('%Y-%m-%d %H:%M')}")
        row += 2
        
        # Report-specific content
        data = report.cached_data or {}
        
        if report.report_type == ReportType.INCOME_STATEMENT.value:
            row = self._excel_income_statement(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        elif report.report_type == ReportType.BALANCE_SHEET.value:
            row = self._excel_balance_sheet(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        elif report.report_type == ReportType.GENERAL_LEDGER.value:
            row = self._excel_general_ledger(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        elif report.report_type == ReportType.SUB_LEDGER.value:
            row = self._excel_sub_ledger(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        elif report.report_type == ReportType.TRIAL_BALANCE.value:
            row = self._excel_trial_balance(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        elif report.report_type == ReportType.EXPENSE_REPORT.value:
            row = self._excel_expense_report(ws, data, row, subheader_font, money_format, thin_border, header_fill, Font)
        
        # Auto-adjust column widths
        for col in range(1, ws.max_column + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
        
        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        content = output.getvalue()
        return content, len(content)
    
    def _excel_income_statement(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write income statement to Excel worksheet"""
        # Revenue section
        ws.cell(row=row, column=1, value='REVENUE').font = subheader_font
        row += 1
        for item in data.get('revenue', []):
            ws.cell(row=row, column=1, value=item.get('account_code', ''))
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('current_amount', 0))
            cell.number_format = money_format
            row += 1
        
        ws.cell(row=row, column=2, value='Total Revenue').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_revenue', 0))
        cell.number_format = money_format
        cell.font = subheader_font
        row += 2
        
        # Cost of Goods
        ws.cell(row=row, column=1, value='COST OF GOODS SOLD').font = subheader_font
        row += 1
        for item in data.get('cost_of_goods', []):
            ws.cell(row=row, column=1, value=item.get('account_code', ''))
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('current_amount', 0))
            cell.number_format = money_format
            row += 1
        
        ws.cell(row=row, column=2, value='Total COGS').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_cost_of_goods', 0))
        cell.number_format = money_format
        row += 1
        
        ws.cell(row=row, column=2, value='GROSS PROFIT').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('gross_profit', 0))
        cell.number_format = money_format
        cell.font = subheader_font
        row += 2
        
        # Operating Expenses
        ws.cell(row=row, column=1, value='OPERATING EXPENSES').font = subheader_font
        row += 1
        for item in data.get('operating_expenses', []):
            ws.cell(row=row, column=1, value=item.get('account_code', ''))
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('current_amount', 0))
            cell.number_format = money_format
            row += 1
        
        ws.cell(row=row, column=2, value='Total Operating Expenses').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_operating_expenses', 0))
        cell.number_format = money_format
        row += 2
        
        # Net Income
        ws.cell(row=row, column=2, value='NET INCOME').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=3, value=data.get('net_income', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        
        return row + 2
    
    def _excel_balance_sheet(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write balance sheet to Excel worksheet"""
        # Assets
        ws.cell(row=row, column=1, value='ASSETS').font = subheader_font
        row += 1
        
        ws.cell(row=row, column=1, value='Current Assets')
        row += 1
        for item in data.get('current_assets', []):
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('balance', 0))
            cell.number_format = money_format
            row += 1
        ws.cell(row=row, column=2, value='Total Current Assets').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_current_assets', 0))
        cell.number_format = money_format
        row += 2
        
        ws.cell(row=row, column=1, value='Fixed Assets')
        row += 1
        for item in data.get('fixed_assets', []):
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('balance', 0))
            cell.number_format = money_format
            row += 1
        ws.cell(row=row, column=2, value='Total Fixed Assets').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_fixed_assets', 0))
        cell.number_format = money_format
        row += 2
        
        ws.cell(row=row, column=2, value='TOTAL ASSETS').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=3, value=data.get('total_assets', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        row += 3
        
        # Liabilities
        ws.cell(row=row, column=1, value='LIABILITIES').font = subheader_font
        row += 1
        for item in data.get('current_liabilities', []):
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('balance', 0))
            cell.number_format = money_format
            row += 1
        ws.cell(row=row, column=2, value='Total Liabilities').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_liabilities', 0))
        cell.number_format = money_format
        row += 2
        
        # Equity
        ws.cell(row=row, column=1, value='EQUITY').font = subheader_font
        row += 1
        for item in data.get('equity', []):
            ws.cell(row=row, column=2, value=item.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=item.get('balance', 0))
            cell.number_format = money_format
            row += 1
        ws.cell(row=row, column=2, value='Retained Earnings')
        cell = ws.cell(row=row, column=3, value=data.get('retained_earnings', 0))
        cell.number_format = money_format
        row += 1
        ws.cell(row=row, column=2, value='Total Equity').font = subheader_font
        cell = ws.cell(row=row, column=3, value=data.get('total_equity', 0))
        cell.number_format = money_format
        row += 2
        
        ws.cell(row=row, column=2, value='TOTAL LIABILITIES & EQUITY').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=3, value=data.get('total_liabilities_and_equity', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        
        return row + 2
    
    def _excel_general_ledger(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write general ledger to Excel worksheet"""
        for account in data.get('accounts', []):
            # Account header
            ws.cell(row=row, column=1, value=f"{account['account_code']} - {account['account_name']}").font = subheader_font
            row += 1
            ws.cell(row=row, column=1, value=f"Opening Balance: {account['opening_balance']:,.2f}")
            row += 1
            
            # Column headers
            headers = ['Date', 'Entry #', 'Description', 'Debit', 'Credit', 'Balance']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = subheader_font
                cell.fill = header_fill
            row += 1
            
            # Entries
            for entry in account.get('entries', []):
                ws.cell(row=row, column=1, value=entry.get('date', ''))
                ws.cell(row=row, column=2, value=entry.get('entry_number', ''))
                ws.cell(row=row, column=3, value=entry.get('description', ''))
                cell = ws.cell(row=row, column=4, value=entry.get('debit', 0))
                cell.number_format = money_format
                cell = ws.cell(row=row, column=5, value=entry.get('credit', 0))
                cell.number_format = money_format
                cell = ws.cell(row=row, column=6, value=entry.get('balance', 0))
                cell.number_format = money_format
                row += 1
            
            # Account totals
            ws.cell(row=row, column=3, value='Account Totals:').font = subheader_font
            cell = ws.cell(row=row, column=4, value=account.get('total_debits', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            cell = ws.cell(row=row, column=5, value=account.get('total_credits', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            cell = ws.cell(row=row, column=6, value=account.get('closing_balance', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            row += 2
        
        # Grand totals
        ws.cell(row=row, column=3, value='GRAND TOTALS:').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=4, value=data.get('total_debits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=5, value=data.get('total_credits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        
        return row + 2
    
    def _excel_trial_balance(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write trial balance to Excel worksheet"""
        # Headers
        headers = ['Account Code', 'Account Name', 'Debit', 'Credit']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = subheader_font
            cell.fill = header_fill
        row += 1
        
        # Lines
        for line in data.get('lines', []):
            ws.cell(row=row, column=1, value=line.get('account_code', ''))
            ws.cell(row=row, column=2, value=line.get('account_name', ''))
            cell = ws.cell(row=row, column=3, value=line.get('debit', 0))
            cell.number_format = money_format
            cell = ws.cell(row=row, column=4, value=line.get('credit', 0))
            cell.number_format = money_format
            row += 1
        
        # Totals
        row += 1
        ws.cell(row=row, column=2, value='TOTALS').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=3, value=data.get('total_debits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=4, value=data.get('total_credits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        row += 1
        
        # Balance check
        is_balanced = data.get('is_balanced', False)
        ws.cell(row=row, column=2, value='Balanced:')
        ws.cell(row=row, column=3, value='Yes' if is_balanced else 'No')
        
        return row + 2
    
    def _excel_sub_ledger(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write sub-ledger (by contact) to Excel worksheet"""
        # Summary info
        ws.cell(row=row, column=1, value=f"Ledger Type: {data.get('ledger_type', 'all').title()}")
        row += 1
        ws.cell(row=row, column=1, value=f"Total Contacts: {data.get('contact_count', 0)}")
        ws.cell(row=row, column=3, value=f"Total Entries: {data.get('entry_count', 0)}")
        row += 2
        
        for contact in data.get('contacts', []):
            # Contact header
            contact_header = f"{contact['contact_name']} ({contact['contact_type']})"
            if contact.get('linked_account_code'):
                contact_header += f" - {contact['linked_account_code']}"
            ws.cell(row=row, column=1, value=contact_header).font = subheader_font
            row += 1
            ws.cell(row=row, column=1, value=f"Opening Balance: {contact['opening_balance']:,.2f}")
            row += 1
            
            # Column headers
            headers = ['Date', 'Entry #', 'Description', 'Reference', 'Debit', 'Credit', 'Balance']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = subheader_font
                cell.fill = header_fill
            row += 1
            
            # Entries
            for entry in contact.get('entries', []):
                ws.cell(row=row, column=1, value=entry.get('date', ''))
                ws.cell(row=row, column=2, value=entry.get('entry_number', ''))
                ws.cell(row=row, column=3, value=entry.get('description', ''))
                ws.cell(row=row, column=4, value=entry.get('reference', ''))
                cell = ws.cell(row=row, column=5, value=entry.get('debit', 0))
                cell.number_format = money_format
                cell = ws.cell(row=row, column=6, value=entry.get('credit', 0))
                cell.number_format = money_format
                cell = ws.cell(row=row, column=7, value=entry.get('balance', 0))
                cell.number_format = money_format
                row += 1
            
            # Contact totals
            ws.cell(row=row, column=3, value='Totals:').font = subheader_font
            cell = ws.cell(row=row, column=5, value=contact.get('total_debits', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            cell = ws.cell(row=row, column=6, value=contact.get('total_credits', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            cell = ws.cell(row=row, column=7, value=contact.get('closing_balance', 0))
            cell.number_format = money_format
            cell.font = subheader_font
            row += 2
        
        # Grand totals
        ws.cell(row=row, column=3, value='GRAND TOTALS:').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=5, value=data.get('total_debits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=6, value=data.get('total_credits', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        
        return row + 2
    
    def _excel_expense_report(self, ws, data, row, subheader_font, money_format, border, header_fill, Font):
        """Write expense report to Excel worksheet"""
        # Headers
        headers = ['Date', 'Description', 'Vendor', 'Category', 'Project', 'Amount']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = subheader_font
            cell.fill = header_fill
        row += 1
        
        # Expenses
        for expense in data.get('expenses', []):
            ws.cell(row=row, column=1, value=expense.get('date', ''))
            ws.cell(row=row, column=2, value=expense.get('description', ''))
            ws.cell(row=row, column=3, value=expense.get('vendor', ''))
            ws.cell(row=row, column=4, value=expense.get('category', ''))
            ws.cell(row=row, column=5, value=expense.get('project', ''))
            cell = ws.cell(row=row, column=6, value=expense.get('amount', 0))
            cell.number_format = money_format
            row += 1
        
        # Total
        row += 1
        ws.cell(row=row, column=5, value='TOTAL').font = Font(bold=True, size=12)
        cell = ws.cell(row=row, column=6, value=data.get('total_amount', 0))
        cell.number_format = money_format
        cell.font = Font(bold=True, size=12)
        
        return row + 2
    
    # =================================================================
    # Word Export
    # =================================================================
    
    def _generate_word(self, report: Report, config: Dict) -> Tuple[bytes, int]:
        """Generate Word document from report data"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(report.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Report Type: {report.report_type.replace('_', ' ').title()}")
        doc.add_paragraph(f"Period: {report.period_display}")
        doc.add_paragraph(f"Generated: {report.generation_completed_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Report-specific content
        data = report.cached_data or {}
        
        if report.report_type == ReportType.INCOME_STATEMENT.value:
            self._word_income_statement(doc, data)
        elif report.report_type == ReportType.BALANCE_SHEET.value:
            self._word_balance_sheet(doc, data)
        elif report.report_type == ReportType.GENERAL_LEDGER.value:
            self._word_general_ledger(doc, data)
        elif report.report_type == ReportType.SUB_LEDGER.value:
            self._word_sub_ledger(doc, data)
        elif report.report_type == ReportType.TRIAL_BALANCE.value:
            self._word_trial_balance(doc, data)
        elif report.report_type == ReportType.EXPENSE_REPORT.value:
            self._word_expense_report(doc, data)
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        content = output.getvalue()
        return content, len(content)
    
    def _word_income_statement(self, doc, data):
        """Write income statement to Word document"""
        doc.add_heading('Income Statement', level=1)
        
        # Revenue section
        doc.add_heading('Revenue', level=2)
        if data.get('revenue'):
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Code'
            hdr_cells[1].text = 'Account'
            hdr_cells[2].text = 'Amount'
            
            for item in data['revenue']:
                row = table.add_row().cells
                row[0].text = item.get('account_code', '')
                row[1].text = item.get('account_name', '')
                row[2].text = f"{item.get('current_amount', 0):,.2f}"
        
        doc.add_paragraph(f"Total Revenue: {data.get('total_revenue', 0):,.2f}")
        doc.add_paragraph(f"Gross Profit: {data.get('gross_profit', 0):,.2f}")
        doc.add_paragraph(f"Operating Income: {data.get('operating_income', 0):,.2f}")
        
        # Summary
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(f"Net Income: {data.get('net_income', 0):,.2f}")
    
    def _word_balance_sheet(self, doc, data):
        """Write balance sheet to Word document"""
        doc.add_heading('Balance Sheet', level=1)
        
        doc.add_heading('Assets', level=2)
        doc.add_paragraph(f"Total Current Assets: {data.get('total_current_assets', 0):,.2f}")
        doc.add_paragraph(f"Total Fixed Assets: {data.get('total_fixed_assets', 0):,.2f}")
        doc.add_paragraph(f"Total Assets: {data.get('total_assets', 0):,.2f}")
        
        doc.add_heading('Liabilities', level=2)
        doc.add_paragraph(f"Total Liabilities: {data.get('total_liabilities', 0):,.2f}")
        
        doc.add_heading('Equity', level=2)
        doc.add_paragraph(f"Total Equity: {data.get('total_equity', 0):,.2f}")
        doc.add_paragraph(f"Total Liabilities & Equity: {data.get('total_liabilities_and_equity', 0):,.2f}")
    
    def _word_general_ledger(self, doc, data):
        """Write general ledger to Word document"""
        doc.add_heading('General Ledger', level=1)
        
        for account in data.get('accounts', []):
            doc.add_heading(f"{account['account_code']} - {account['account_name']}", level=2)
            doc.add_paragraph(f"Opening Balance: {account['opening_balance']:,.2f}")
            
            if account.get('entries'):
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ['Date', 'Entry #', 'Description', 'Debit', 'Credit']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for entry in account['entries']:
                    row = table.add_row().cells
                    row[0].text = entry.get('date', '')
                    row[1].text = entry.get('entry_number', '')
                    row[2].text = entry.get('description', '')[:50]
                    row[3].text = f"{entry.get('debit', 0):,.2f}"
                    row[4].text = f"{entry.get('credit', 0):,.2f}"
            
            doc.add_paragraph(f"Closing Balance: {account['closing_balance']:,.2f}")
    
    def _word_sub_ledger(self, doc, data):
        """Write sub-ledger (by contact) to Word document"""
        doc.add_heading('Sub-Ledger Report', level=1)
        
        doc.add_paragraph(f"Ledger Type: {data.get('ledger_type', 'all').title()}")
        doc.add_paragraph(f"Total Contacts: {data.get('contact_count', 0)}")
        doc.add_paragraph(f"Total Entries: {data.get('entry_count', 0)}")
        doc.add_paragraph()
        
        for contact in data.get('contacts', []):
            contact_header = f"{contact['contact_name']} ({contact['contact_type']})"
            if contact.get('linked_account_code'):
                contact_header += f" - {contact['linked_account_code']}"
            doc.add_heading(contact_header, level=2)
            doc.add_paragraph(f"Opening Balance: {contact['opening_balance']:,.2f}")
            
            if contact.get('entries'):
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ['Date', 'Entry #', 'Description', 'Debit', 'Credit', 'Balance']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for entry in contact['entries'][:20]:  # Limit to 20 entries per contact in Word
                    row = table.add_row().cells
                    row[0].text = entry.get('date', '')
                    row[1].text = entry.get('entry_number', '')
                    row[2].text = entry.get('description', '')[:40]
                    row[3].text = f"{entry.get('debit', 0):,.2f}"
                    row[4].text = f"{entry.get('credit', 0):,.2f}"
                    row[5].text = f"{entry.get('balance', 0):,.2f}"
                
                if len(contact['entries']) > 20:
                    doc.add_paragraph(f"... and {len(contact['entries']) - 20} more entries")
            
            doc.add_paragraph(f"Closing Balance: {contact['closing_balance']:,.2f}")
            doc.add_paragraph()
        
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(f"Total Debits: {data.get('total_debits', 0):,.2f}")
        doc.add_paragraph(f"Total Credits: {data.get('total_credits', 0):,.2f}")
    
    def _word_trial_balance(self, doc, data):
        """Write trial balance to Word document"""
        doc.add_heading('Trial Balance', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Code', 'Account', 'Debit', 'Credit']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for line in data.get('lines', []):
            row = table.add_row().cells
            row[0].text = line.get('account_code', '')
            row[1].text = line.get('account_name', '')
            row[2].text = f"{line.get('debit', 0):,.2f}"
            row[3].text = f"{line.get('credit', 0):,.2f}"
        
        doc.add_paragraph()
        doc.add_paragraph(f"Total Debits: {data.get('total_debits', 0):,.2f}")
        doc.add_paragraph(f"Total Credits: {data.get('total_credits', 0):,.2f}")
        doc.add_paragraph(f"Balanced: {'Yes' if data.get('is_balanced') else 'No'}")
    
    def _word_expense_report(self, doc, data):
        """Write expense report to Word document"""
        doc.add_heading('Expense Report', level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Date', 'Description', 'Vendor', 'Category', 'Amount']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for expense in data.get('expenses', []):
            row = table.add_row().cells
            row[0].text = expense.get('date', '')
            row[1].text = expense.get('description', '')[:30]
            row[2].text = expense.get('vendor', '')
            row[3].text = expense.get('category', '')
            row[4].text = f"{expense.get('amount', 0):,.2f}"
        
        doc.add_paragraph()
        doc.add_paragraph(f"Total Amount: {data.get('total_amount', 0):,.2f}")
    
    # =================================================================
    # CSV Export
    # =================================================================
    
    def _generate_csv(self, report: Report, config: Dict) -> Tuple[bytes, int]:
        """Generate CSV from report data"""
        import csv
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header info
        writer.writerow([report.name])
        writer.writerow([f"Period: {report.period_display}"])
        writer.writerow([])
        
        data = report.cached_data or {}
        
        if report.report_type == ReportType.TRIAL_BALANCE.value:
            writer.writerow(['Account Code', 'Account Name', 'Debit', 'Credit'])
            for line in data.get('lines', []):
                writer.writerow([
                    line.get('account_code'),
                    line.get('account_name'),
                    line.get('debit'),
                    line.get('credit')
                ])
            writer.writerow([])
            writer.writerow(['', 'TOTALS', data.get('total_debits'), data.get('total_credits')])
        
        elif report.report_type == ReportType.EXPENSE_REPORT.value:
            writer.writerow(['Date', 'Description', 'Vendor', 'Category', 'Project', 'Amount'])
            for expense in data.get('expenses', []):
                writer.writerow([
                    expense.get('date'),
                    expense.get('description'),
                    expense.get('vendor'),
                    expense.get('category'),
                    expense.get('project'),
                    expense.get('amount')
                ])
            writer.writerow([])
            writer.writerow(['', '', '', '', 'TOTAL', data.get('total_amount')])
        
        content = output.getvalue().encode('utf-8')
        return content, len(content)
    
    # =================================================================
    # JSON Export
    # =================================================================
    
    def _generate_json(self, report: Report, config: Dict) -> Tuple[bytes, int]:
        """Generate JSON from report data"""
        import json
        
        export_data = {
            'report': {
                'number': report.report_number,
                'name': report.name,
                'type': report.report_type,
                'period_start': report.period_start.isoformat(),
                'period_end': report.period_end.isoformat(),
                'generated_at': report.generation_completed_at.isoformat() if report.generation_completed_at else None,
            },
            'summary': report.summary_totals,
            'data': report.cached_data
        }
        
        content = json.dumps(export_data, indent=2, default=str).encode('utf-8')
        return content, len(content)
    
    # =================================================================
    # PDF Export (placeholder)
    # =================================================================
    
    def _generate_pdf(self, report: Report, config: Dict) -> Tuple[bytes, int]:
        """Generate PDF from report data"""
        # PDF generation requires additional libraries like reportlab or weasyprint
        # For now, return a placeholder or convert from HTML
        raise NotImplementedError(
            "PDF export requires additional setup. "
            "Please use Excel or Word export, or install reportlab: pip install reportlab"
        )
